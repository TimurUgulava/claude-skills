#!/usr/bin/env python3
"""
export_docx.py — экспорт MD-артефакта Прохора в .docx в фирменном стиле.

Универсальный для всех этапов (паспорт, ЦА, концепция, тексты): обложка с заголовком
и мета-таблицей + рендер тела MD + footer.

Использование:
    python3 export_docx.py passport.md --output passport.docx \
        --title "МАРКЕТИНГОВЫЙ ПАСПОРТ ОБЪЕКТА" \
        --subtitle "ЖК «Пример» · комфорт-класс · Санкт-Петербург" \
        --meta-json '{"Объект":"ЖК «Пример»","Застройщик":"<застройщик>"}' \
        --footer "ваш бренд / канал"

Стиль: оранжевый акцент #FE9901, Calibri, зебра в таблицах. Подпись в footer
настраивается флагом --footer (по умолчанию — без подписи бренда).
"""

import argparse
import json
import re
import sys
from datetime import date
from pathlib import Path

try:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Cm, Pt, RGBColor
except ImportError:
    print("ERROR: python-docx не установлен. Установи: pip install python-docx", file=sys.stderr)
    sys.exit(1)


ORANGE = RGBColor(0xFE, 0x99, 0x01)
DARK = RGBColor(0x1E, 0x1E, 0x1E)
GREY = RGBColor(0x70, 0x70, 0x70)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
ZEBRA = "F7F7F7"
ORANGE_HEX = "FE9901"


def shade_cell(cell, fill_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def add_heading(doc, text, color=DARK, size=14, space_before=18, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = "Calibri"
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_cover_meta(doc, meta: dict):
    rows = list(meta.items())
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (k, v) in enumerate(rows):
        cell_k = table.rows[i].cells[0]
        cell_v = table.rows[i].cells[1]
        cell_k.text = ""
        cell_v.text = ""
        run_k = cell_k.paragraphs[0].add_run(str(k))
        run_k.bold = True
        run_k.font.size = Pt(10)
        run_k.font.color.rgb = ORANGE
        run_v = cell_v.paragraphs[0].add_run(str(v))
        run_v.font.size = Pt(10)
        run_v.font.color.rgb = DARK
        if i % 2 == 1:
            shade_cell(cell_k, ZEBRA)
            shade_cell(cell_v, ZEBRA)


def parse_md_table(lines, start):
    cells = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        if re.match(r"^\|[-:\s|]+\|$", lines[i].strip()):
            i += 1
            continue
        row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        cells.append(row)
        i += 1
    return cells, i


def render_md_table(doc, cells):
    if not cells:
        return
    cols = max(len(r) for r in cells)
    table = doc.add_table(rows=len(cells), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, row in enumerate(cells):
        for j in range(cols):
            cell = table.rows[i].cells[j]
            cell.text = ""
            text = row[j] if j < len(row) else ""
            text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
            text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
            run = cell.paragraphs[0].add_run(text)
            run.font.size = Pt(10)
            if i == 0:
                run.bold = True
                run.font.color.rgb = WHITE
                shade_cell(cell, ORANGE_HEX)
            elif i % 2 == 0:
                shade_cell(cell, ZEBRA)


def add_inline_runs(paragraph, text, base_size=11):
    parts = re.split(r"(\*\*[^*]+\*\*|\[[^\]]+\]\([^)]+\))", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(base_size)
            run.font.color.rgb = DARK
        elif re.match(r"\[[^\]]+\]\([^)]+\)", part):
            m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", part)
            run = paragraph.add_run(m.group(1))
            run.font.size = Pt(base_size)
            run.font.color.rgb = ORANGE
            run.underline = True
        else:
            run = paragraph.add_run(part)
            run.font.size = Pt(base_size)
            run.font.color.rgb = DARK


def add_footer(doc, role, brand=""):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    parts = [brand, role, date.today().isoformat()]
    text = " · ".join(part for part in parts if part)
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = GREY


def render_md_to_docx(md_text, doc):
    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        if stripped.startswith("# "):
            add_heading(doc, stripped[2:].strip(), color=ORANGE, size=18, space_before=12)
            i += 1
        elif stripped.startswith("## "):
            add_heading(doc, stripped[3:].strip(), color=DARK, size=14, space_before=18)
            i += 1
        elif stripped.startswith("### "):
            add_heading(doc, stripped[4:].strip(), color=ORANGE, size=12, space_before=10)
            i += 1
        elif stripped.startswith("|"):
            cells, i = parse_md_table(lines, i)
            render_md_table(doc, cells)
        elif re.match(r"^\d+\.\s+", stripped):
            while i < len(lines) and re.match(r"^\d+\.\s+", lines[i].strip()):
                p = doc.add_paragraph(style="List Number")
                text = re.sub(r"^\d+\.\s+", "", lines[i].strip())
                add_inline_runs(p, text)
                i += 1
        elif stripped.startswith("- ") or stripped.startswith("* "):
            while i < len(lines) and (lines[i].strip().startswith("- ") or lines[i].strip().startswith("* ")):
                p = doc.add_paragraph(style="List Bullet")
                text = lines[i].strip().lstrip("-*").strip()
                add_inline_runs(p, text)
                i += 1
        elif stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.7)
            run = p.add_run(stripped[2:].strip())
            run.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = GREY
            i += 1
        elif re.match(r"^!\[[^\]]*\]\(([^)]+)\)\s*$", stripped):
            m = re.match(r"^!\[[^\]]*\]\(([^)]+)\)\s*$", stripped)
            img_path = m.group(1).strip()
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(10)
            try:
                p.add_run().add_picture(img_path, width=Cm(13))
            except Exception:
                run = p.add_run(f"[изображение не найдено: {img_path}]")
                run.font.size = Pt(9)
                run.font.color.rgb = GREY
            i += 1
        else:
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.3
            p.paragraph_format.space_after = Pt(8)
            add_inline_runs(p, stripped)
            i += 1


def strip_frontmatter(text):
    if text.startswith("---\n"):
        end = text.find("\n---\n", 4)
        if end != -1:
            return text[end + 5:]
    return text


def main():
    parser = argparse.ArgumentParser(description="Экспорт MD → .docx в фирменном стиле Прохора")
    parser.add_argument("md", help="Путь к MD-артефакту")
    parser.add_argument("--output", required=True, help="Путь к выходному .docx")
    parser.add_argument("--title", default="МАРКЕТИНГОВЫЙ ПАСПОРТ ОБЪЕКТА", help="Заголовок обложки")
    parser.add_argument("--subtitle", default="", help="Подзаголовок обложки")
    parser.add_argument("--meta-json", default="{}", help="JSON-объект для мета-таблицы обложки")
    parser.add_argument("--role", default="подготовлено Прохором", help="Строка роли в footer")
    parser.add_argument("--footer", default="", help="Подпись бренда/канала в footer (по умолчанию пусто)")
    args = parser.parse_args()

    md_text = strip_frontmatter(Path(args.md).read_text(encoding="utf-8"))
    meta = json.loads(args.meta_json)
    meta.setdefault("Дата", date.today().isoformat())
    meta.setdefault("Подготовил", "Прохор · AI-маркетолог девелопера")

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    add_heading(doc, args.title, color=ORANGE, size=20, space_before=0,
                align=WD_PARAGRAPH_ALIGNMENT.LEFT)
    if args.subtitle:
        sp = doc.add_paragraph()
        srun = sp.add_run(args.subtitle)
        srun.font.size = Pt(12)
        srun.font.color.rgb = DARK
        srun.bold = True
        sp.paragraph_format.space_after = Pt(10)
    add_cover_meta(doc, meta)
    doc.add_paragraph()

    render_md_to_docx(md_text, doc)
    add_footer(doc, args.role, args.footer)

    Path(args.output).parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.output)
    print(f"OK: docx saved to {args.output}", file=sys.stderr)


if __name__ == "__main__":
    main()
